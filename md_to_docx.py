"""
md_to_docx.py  —  Convert report/architecture.md to report/architecture.docx
Usage: python md_to_docx.py
"""

import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MD_PATH  = Path(__file__).parent / "report" / "architecture.md"
OUT_PATH = Path(__file__).parent / "report" / "architecture.docx"


def add_shading(paragraph, fill="F2F2F2"):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill)
    pPr.append(shd)


def add_hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "AAAAAA")
    pBdr.append(bot)
    pPr.append(pBdr)


def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    add_shading(p)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)


def inline_format(para, text):
    """Render **bold** and `code` inline spans."""
    TICK = "`"
    pattern = r"(\*\*[^*]+\*\*|" + TICK + r"[^" + TICK + r"]+" + TICK + r")"
    parts = re.split(pattern, text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith(TICK) and part.endswith(TICK):
            run = para.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        else:
            para.add_run(part)


def flush_table(doc, rows):
    if not rows:
        return
    # Remove pure-separator rows (---|---) but keep data rows
    data_rows = [r for r in rows if not all(
        re.match(r"^[-:\s]+$", c.strip()) for c in r if c.strip()
    )]
    if not data_rows:
        return
    ncols = max(len(r) for r in data_rows)
    tbl = doc.add_table(rows=len(data_rows), cols=ncols)
    tbl.style = "Table Grid"
    for ri, row in enumerate(data_rows):
        for ci in range(ncols):
            cell_text = row[ci].strip() if ci < len(row) else ""
            cell = tbl.rows[ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            if ri == 0:
                run = p.add_run(cell_text)
                run.bold = True
            else:
                inline_format(p, cell_text)
    doc.add_paragraph()


def convert(md_path: Path, out_path: Path):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # Default font
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    lines = md_path.read_text(encoding="utf-8").splitlines()

    in_code  = False
    in_table = False
    code_buf = []
    table_rows: list = []

    for raw in lines:
        line    = raw.rstrip()
        stripped = line.strip()

        # ── code fence ────────────────────────────────────────
        if stripped.startswith("```"):
            if not in_code:
                in_code  = True
                code_buf = []
            else:
                in_code = False
                add_code_block(doc, "\n".join(code_buf))
            continue
        if in_code:
            code_buf.append(line)
            continue

        # ── table row ─────────────────────────────────────────
        if stripped.startswith("|") and "|" in stripped[1:]:
            in_table = True
            cells = [c for c in stripped.split("|")]
            if cells and not cells[0].strip():  cells = cells[1:]
            if cells and not cells[-1].strip(): cells = cells[:-1]
            table_rows.append(cells)
            continue
        else:
            if in_table:
                in_table = False
                flush_table(doc, table_rows)
                table_rows = []

        # ── image  ![caption](path) ───────────────────────────
        m_img = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)", stripped)
        if m_img:
            caption  = m_img.group(1)
            img_path = Path(m_img.group(2))
            # Resolve relative to the MD file's directory
            if not img_path.is_absolute():
                img_path = md_path.parent / img_path
            if img_path.exists():
                doc.add_picture(str(img_path), width=Inches(5.5))
                last_para = doc.paragraphs[-1]
                last_para.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER
                if caption:
                    cap = doc.add_paragraph(caption)
                    cap.alignment = 1
                    cap.paragraph_format.space_before = Pt(2)
                    cap.paragraph_format.space_after  = Pt(10)
                    cap.runs[0].font.size   = Pt(9)
                    cap.runs[0].font.italic = True
            else:
                p = doc.add_paragraph()
                run = p.add_run(f"[Figure not found: {img_path.name}]")
                run.font.italic = True
                run.font.color.rgb = RGBColor(0xAA, 0x00, 0x00)
            continue

        # ── heading ───────────────────────────────────────────
        m = re.match(r"^(#{1,4})\s+(.*)", stripped)
        if m:
            level = len(m.group(1))
            h = doc.add_heading(m.group(2), level=level)
            h.paragraph_format.space_before = Pt(10)
            continue

        # ── horizontal rule ───────────────────────────────────
        if re.match(r"^[-*_]{3,}$", stripped):
            add_hr(doc)
            continue

        # ── blank line ────────────────────────────────────────
        if not stripped:
            continue

        # ── bullet ───────────────────────────────────────────
        m = re.match(r"^[-*]\s+(.*)", stripped)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            inline_format(p, m.group(1))
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            continue

        # ── numbered list ─────────────────────────────────────
        m = re.match(r"^\d+\.\s+(.*)", stripped)
        if m:
            p = doc.add_paragraph(style="List Number")
            inline_format(p, m.group(1))
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            continue

        # ── normal paragraph ──────────────────────────────────
        p = doc.add_paragraph()
        inline_format(p, stripped)
        p.paragraph_format.space_after = Pt(6)

    # flush any open table
    if in_table:
        flush_table(doc, table_rows)

    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    convert(MD_PATH, OUT_PATH)
